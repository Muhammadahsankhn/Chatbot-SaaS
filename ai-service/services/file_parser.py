"""
ai-service/services/file_parser.py

Parses uploaded files into plain text.
Supported: PDF, Word (.docx), Excel (.xlsx/.xls), CSV, TXT, JSON

Install dependencies:
  pip install pdfplumber python-docx openpyxl
"""

import csv
import io
import json
import os
from typing import List, Tuple

# ── Chunk settings ──
CHUNK_SIZE    = 500   # words per chunk
CHUNK_OVERLAP = 50    # words overlap between chunks


# ──────────────────────────────────────────────
# CHUNK TEXT
# ──────────────────────────────────────────────
def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    """Split plain text into overlapping word chunks."""
    words  = text.split()
    chunks = []
    i      = 0

    while i < len(words):
        chunk = words[i : i + chunk_size]
        chunks.append(" ".join(chunk))
        i += chunk_size - overlap

    return [c for c in chunks if len(c.strip()) > 30]


# ──────────────────────────────────────────────
# PDF
# ──────────────────────────────────────────────
def parse_pdf(file_bytes: bytes) -> str:
    try:
        import pdfplumber
        text_parts = []

        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                # Extract regular text
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

                # Extract tables as text
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        clean_row = [str(cell).strip() for cell in row if cell and str(cell).strip()]
                        if clean_row:
                            text_parts.append(" | ".join(clean_row))

        result = "\n".join(text_parts)
        print(f"[Parser] PDF → {len(result)} chars, {len(result.split())} words")
        return result

    except ImportError:
        raise ImportError("pdfplumber not installed. Run: pip install pdfplumber")
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {e}")


# ──────────────────────────────────────────────
# WORD (.docx)
# ──────────────────────────────────────────────
def parse_word(file_bytes: bytes) -> str:
    try:
        from docx import Document

        doc        = Document(io.BytesIO(file_bytes))
        text_parts = []

        # Paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_parts.append(text)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    text_parts.append(" | ".join(cells))

        result = "\n".join(text_parts)
        print(f"[Parser] DOCX → {len(result)} chars, {len(result.split())} words")
        return result

    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")
    except Exception as e:
        raise ValueError(f"Failed to parse Word document: {e}")


# ──────────────────────────────────────────────
# EXCEL (.xlsx / .xls)
# ──────────────────────────────────────────────
def parse_excel(file_bytes: bytes) -> str:
    try:
        import openpyxl

        wb         = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
        text_parts = []

        for sheet in wb.worksheets:
            text_parts.append(f"Sheet: {sheet.title}")

            rows       = list(sheet.iter_rows(values_only=True))
            headers    = []

            for i, row in enumerate(rows):
                clean = [str(cell).strip() for cell in row if cell is not None and str(cell).strip() not in ("", "None")]
                if not clean:
                    continue

                # First non-empty row = headers
                if not headers:
                    headers = clean
                    text_parts.append("Columns: " + " | ".join(headers))
                    continue

                # Format data rows as "Header: Value" pairs
                if len(clean) == len(headers):
                    pairs = [f"{headers[j]}: {clean[j]}" for j in range(len(headers))]
                    text_parts.append(", ".join(pairs))
                else:
                    text_parts.append(" | ".join(clean))

        result = "\n".join(text_parts)
        print(f"[Parser] XLSX → {len(result)} chars, {len(result.split())} words")
        return result

    except ImportError:
        raise ImportError("openpyxl not installed. Run: pip install openpyxl")
    except Exception as e:
        raise ValueError(f"Failed to parse Excel file: {e}")


# ──────────────────────────────────────────────
# CSV
# ──────────────────────────────────────────────
def parse_csv(file_bytes: bytes) -> str:
    try:
        # Try UTF-8 first, fallback to latin-1
        try:
            content = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            content = file_bytes.decode("latin-1")

        reader     = csv.reader(io.StringIO(content))
        rows       = list(reader)
        text_parts = []
        headers    = []

        for i, row in enumerate(rows):
            clean = [cell.strip() for cell in row if cell.strip()]
            if not clean:
                continue

            if i == 0:
                headers = clean
                text_parts.append("Columns: " + " | ".join(headers))
                continue

            if headers and len(clean) == len(headers):
                pairs = [f"{headers[j]}: {clean[j]}" for j in range(len(headers))]
                text_parts.append(", ".join(pairs))
            else:
                text_parts.append(" | ".join(clean))

        result = "\n".join(text_parts)
        print(f"[Parser] CSV → {len(result)} chars, {len(result.split())} words")
        return result

    except Exception as e:
        raise ValueError(f"Failed to parse CSV: {e}")


# ──────────────────────────────────────────────
# TXT
# ──────────────────────────────────────────────
def parse_txt(file_bytes: bytes) -> str:
    try:
        try:
            result = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            result = file_bytes.decode("latin-1")

        print(f"[Parser] TXT → {len(result)} chars, {len(result.split())} words")
        return result

    except Exception as e:
        raise ValueError(f"Failed to parse TXT: {e}")


# ──────────────────────────────────────────────
# JSON
# ──────────────────────────────────────────────
def parse_json(file_bytes: bytes) -> str:
    try:
        data       = json.loads(file_bytes.decode("utf-8"))
        text_parts = []

        def flatten(obj, prefix=""):
            if isinstance(obj, dict):
                for k, v in obj.items():
                    flatten(v, f"{prefix}{k}: " if not prefix else f"{prefix} > {k}: ")
            elif isinstance(obj, list):
                for i, item in enumerate(obj):
                    flatten(item, prefix)
            else:
                val = str(obj).strip()
                if val and val.lower() not in ("none", "null", ""):
                    text_parts.append(f"{prefix}{val}")

        flatten(data)
        result = "\n".join(text_parts)
        print(f"[Parser] JSON → {len(result)} chars, {len(result.split())} words")
        return result

    except Exception as e:
        raise ValueError(f"Failed to parse JSON: {e}")


# ──────────────────────────────────────────────
# MAIN ENTRY — auto-detect file type
# ──────────────────────────────────────────────
def parse_file(filename: str, file_bytes: bytes) -> Tuple[str, List[str]]:
    """
    Parse any supported file.

    Returns:
        (raw_text, chunks)
        raw_text: full extracted text
        chunks:   list of 500-word chunks ready for embedding
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    print(f"[Parser] Parsing '{filename}' (ext={ext}, size={len(file_bytes)} bytes)")

    if ext == "pdf":
        raw_text = parse_pdf(file_bytes)
    elif ext == "docx":
        raw_text = parse_word(file_bytes)
    elif ext in ("xlsx", "xls"):
        raw_text = parse_excel(file_bytes)
    elif ext == "csv":
        raw_text = parse_csv(file_bytes)
    elif ext == "txt":
        raw_text = parse_txt(file_bytes)
    elif ext == "json":
        raw_text = parse_json(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: .{ext}. Supported: PDF, DOCX, XLSX, XLS, CSV, TXT, JSON")

    if not raw_text or len(raw_text.strip()) < 50:
        raise ValueError("File appears to be empty or contains too little text.")

    chunks = chunk_text(raw_text)
    print(f"[Parser] ✅ Done → {len(chunks)} chunks from '{filename}'")

    return raw_text, chunks